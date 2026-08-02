from __future__ import annotations

import base64
from io import BytesIO
from uuid import uuid4
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document
from docx.shared import Inches

from docmind_ai.parsing import DocumentParsingService
from docmind_ai.parsing.common import DocumentParsingError
from docmind_ai.parsing.docx_parser import DocxParser
from docmind_ai.parsing.ocr import OcrTextBlock
from docmind_ai.parsing.ooxml_security import validate_ooxml_archive
from docmind_ai.parsing.pdf_parser import PdfParser
from pdf_fixture import build_pdf

PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def build_docx() -> bytes:
    document = Document()
    document.core_properties.title = "测试合同"
    document.sections[0].header.paragraphs[0].text = "公司页眉"
    document.sections[0].footer.paragraphs[0].text = "第 PAGE 页"
    document.add_heading("合同标题", level=1)
    document.add_paragraph("联系人: alice@example.com, 电话: +86 13800138000")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "值"
    table.cell(1, 0).text = "金额"
    table.cell(1, 1).text = "100 元"
    document.add_picture(BytesIO(PNG_1X1), width=Inches(0.2))
    document.add_page_break()
    document.add_paragraph("第二页")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_docx_parser_preserves_structure_resources_and_locations() -> None:
    response = DocxParser().parse(
        source_version_id=uuid4(),
        source_format="docx",
        content=build_docx(),
        language="zh-CN",
    )

    block_types = [block.type for block in response.document.blocks]
    assert response.document.metadata.title == "测试合同"
    assert "heading" in block_types
    assert "table" in block_types
    assert "image" in block_types
    assert "page_break" in block_types
    assert response.document.headers[0].blocks[0].type == "paragraph"
    assert response.document.footers[0].blocks[0].type == "paragraph"
    assert len(response.resources) == 1
    assert response.resources[0].content_base64 is not None
    assert {node.kind for node in response.text_nodes} >= {
        "heading",
        "paragraph",
        "table_cell",
        "header",
        "footer",
    }
    assert any(node.page_number == 2 and node.text == "第二页" for node in response.text_nodes)


def test_ooxml_preflight_rejects_path_traversal() -> None:
    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "content")
        archive.writestr("word/document.xml", "document")
        archive.writestr("../outside.xml", "unsafe")

    with pytest.raises(DocumentParsingError, match="不安全"):
        validate_ooxml_archive(output.getvalue())


def test_pdf_parser_preserves_page_and_bbox() -> None:
    response = PdfParser(max_pages=10).parse(
        source_version_id=uuid4(),
        source_format="pdf",
        content=build_pdf(text="Contract amount 100"),
        language="en-US",
    )

    assert response.document.metadata.source_page_count == 1
    assert response.document.blocks[0].type == "page_marker"
    assert response.text_nodes[0].text == "Contract amount 100"
    assert response.text_nodes[0].page_number == 1
    assert len(response.text_nodes[0].metadata["bbox"]) == 4  # type: ignore[arg-type]


class StubOcrAdapter:
    @property
    def available(self) -> bool:
        return True

    def extract_page(self, *, pdf_content: bytes, page_number: int) -> list[OcrTextBlock]:
        assert pdf_content.startswith(b"%PDF-")
        assert page_number == 1
        return [
            OcrTextBlock(
                text="OCR text",
                left=10,
                top=10,
                right=100,
                bottom=30,
                confidence=0.9,
            )
        ]


def test_pdf_parser_uses_optional_ocr_adapter() -> None:
    response = PdfParser(max_pages=10, ocr_adapter=StubOcrAdapter()).parse(
        source_version_id=uuid4(),
        source_format="pdf",
        content=build_pdf(text=None),
        language="en-US",
    )

    assert response.text_nodes[0].text == "OCR text"
    assert [warning.code for warning in response.warnings] == ["OCR_TEXT_USED"]


def test_doc_requires_configured_libreoffice() -> None:
    service = DocumentParsingService(libreoffice_binary="missing-docmind-libreoffice")

    with pytest.raises(DocumentParsingError, match="未配置"):
        service.parse(
            source_version_id=uuid4(),
            source_format="doc",
            content=bytes.fromhex("D0CF11E0A1B11AE1") + b"legacy",
            language="zh-CN",
        )
