#!/usr/bin/env python3
"""Create the deterministic DOCX fixture used by the real-infrastructure smoke test."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = RGBColor(0x24, 0x2A, 0x33)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x64, 0x6B, 0x75)
TABLE_FILL = "F2F4F7"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
# Named interoperability override: the bundled LibreOffice renderer exposes
# Arial Unicode MS consistently for CJK glyphs, while preserving Calibri for Latin text.
EAST_ASIA_FONT = "Arial Unicode MS"


def set_run_font(
    run,
    *,
    size: float = 11,
    color: RGBColor = INK,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = "Calibri"
    run_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    run_fonts.set(qn("w:ascii"), "Calibri")
    run_fonts.set(qn("w:hAnsi"), "Calibri")
    run_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_width(cell, width: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    tc_width = properties.first_child_found_in("w:tcW")
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        properties.append(tc_width)
    tc_width.set(qn("w:w"), str(width))
    tc_width.set(qn("w:type"), "dxa")


def set_cell_margins(cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in CELL_MARGINS_DXA.items():
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr
    table_width = properties.first_child_found_in("w:tblW")
    table_width.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    properties.append(indent)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = paragraph.add_run("第 ")
    set_run_font(prefix, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=9, color=MUTED)
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(
        qn("w:eastAsia"), EAST_ASIA_FONT
    )

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.element.get_or_add_rPr().get_or_add_rFonts().set(
            qn("w:eastAsia"), EAST_ASIA_FONT
        )


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    left = paragraph.add_run("DOCMIND 联调样例")
    set_run_font(left, size=9, color=MUTED, bold=True)
    right = paragraph.add_run("    内部测试 · 请勿外传")
    set_run_font(right, size=9, color=MUTED)

    footer = section.footer
    add_page_field(footer.paragraphs[0])


def add_metadata(document: Document) -> None:
    rows = [
        ("文档类型", "跨境供应商尽调简报"),
        ("负责人", "林晓岚"),
        ("日期", "2026 年 8 月 1 日"),
        ("状态", "待审核"),
    ]
    for label, value in rows:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}：")
        set_run_font(label_run, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run)


def add_contacts_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["国家/地区", "联系人", "联系信息"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, TABLE_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=10.5, bold=True)

    data = [
        ("中国", "张伟", "+86 138 0013 8000；zhang.wei@example.cn；11010519491231002X"),
        ("美国", "Emily Carter", "+1 (415) 555-0138；emily.carter@example.com"),
        ("日本", "佐藤健", "+81 90-1234-5678；sato.ken@example.jp"),
        ("德国", "Anna Müller", "+49 30 12345678；anna.mueller@example.de"),
    ]
    for country, contact, info in data:
        cells = table.add_row().cells
        for index, value in enumerate((country, contact, info)):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=10)
    set_table_geometry(table, [1500, 2100, 5760])


def add_primary_contact_fields(document: Document) -> None:
    """Add deterministic key/value text for extraction and token-restoration checks."""
    rows = [
        ("联系人姓名", "张伟"),
        ("联系人电话", "+86 138 0013 8000"),
        ("联系人电子邮箱", "zhang.wei@example.cn"),
    ]
    for label, value in rows:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}：")
        set_run_font(label_run, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run)


def build(output: Path) -> None:
    document = Document()
    configure_styles(document)
    configure_section(document.sections[0])

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("跨境供应商尽调简报")
    set_run_font(run, size=23, color=INK, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("用于验证 DOCX 解析、版式复刻、敏感信息占位与模板版本管理")
    set_run_font(run, size=12, color=MUTED)
    add_metadata(document)

    document.add_heading("一、审查摘要", level=1)
    document.add_paragraph(
        "本简报包含标题、页眉页脚、分级标题、混合样式、表格和多国家联系信息，"
        "用于验证系统在真实基础设施下能够保存原件、生成安全 HTML，并保持可编辑结构。"
    )
    paragraph = document.add_paragraph()
    paragraph.add_run("重点结论：").bold = True
    paragraph.add_run(
        "供应商资料完整，但所有个人信息在发送给模型前必须完成脱敏，并在返回后按映射恢复。"
    )

    document.add_heading("二、联系人信息", level=1)
    citation = document.add_paragraph("表 1 · 本地联调虚构数据，仅用于自动化测试")
    citation.paragraph_format.space_before = Pt(4)
    citation.paragraph_format.space_after = Pt(4)
    for run in citation.runs:
        set_run_font(run, size=9, color=MUTED, italic=True)
    add_primary_contact_fields(document)
    add_contacts_table(document)

    document.add_page_break()
    document.add_heading("三、审查记录", level=1)
    document.add_heading("3.1 合规要求", level=2)
    document.add_paragraph(
        "应用层需要识别中国、美国、日本、韩国、德国、法国、英国、澳大利亚和荷兰等地区的常见 PII 格式。"
    )
    document.add_heading("3.2 模板要求", level=2)
    paragraph = document.add_paragraph()
    first = paragraph.add_run("版式复刻")
    set_run_font(first, bold=True, color=BLUE)
    second = paragraph.add_run("应保留页面、段落、字号、对齐、页眉页脚与表格结构；")
    set_run_font(second)
    third = paragraph.add_run("编辑版本")
    set_run_font(third, bold=True, color=BLUE)
    fourth = paragraph.add_run("必须可追溯、可发布并可回滚。")
    set_run_font(fourth)

    document.add_heading("四、审批", level=1)
    table = document.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    values = [
        ("角色", "姓名", "结论"),
        ("合规负责人", "王敏", "通过，待模板复核"),
    ]
    for row_index, row_values in enumerate(values):
        for column_index, value in enumerate(row_values):
            cell = table.rows[row_index].cells[column_index]
            if row_index == 0:
                shade_cell(cell, TABLE_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=10.5, bold=row_index == 0)
    set_table_geometry(table, [2200, 2200, 4960])

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: create_e2e_docx.py OUTPUT.docx")
    build(Path(sys.argv[1]).resolve())


if __name__ == "__main__":
    main()
